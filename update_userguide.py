#!/usr/bin/env python3
"""
Update userguide.docx with enhanced retail security context
"""

from docx import Document
from docx.shared import Pt, RGBColor

def main():
    # Open existing userguide.docx
    doc = Document('userguide.docx')
    
    # Find the overview section and add retail context after it
    found_overview = False
    for i, para in enumerate(doc.paragraphs):
        # Look for overview text to insert after
        if 'WebAI Smart Retail Security is a browser-based' in para.text:
            found_overview = True
            # Skip to after the next paragraph
            continue
        
        if found_overview and 'The system provides instant visual feedback' in para.text:
            # Insert new paragraph after this one
            # Add retail justification
            idx = doc.paragraphs.index(para) + 1
            
            # Insert at position by adding after we find implementation summary
            break
    
    # Just append at the end of document before submission checklist
    insertion_point = None
    for i, para in enumerate(doc.paragraphs):
        if 'IMPLEMENTATION SUMMARY' in para.text:
            insertion_point = i
            break
    
    if insertion_point:
        # Create new paragraph with retail context
        # We'll insert it right after the overview paragraphs
        retail_para = doc.paragraphs[2].insert_paragraph_before()
        retail_para.style = 'Normal'
        
        title_run = retail_para.add_run('\n\nRetail Security Class Justification: ')
        title_run.bold = True
        title_run.font.size = Pt(10)
        
        context_run = retail_para.add_run(
            'The COCO-SSD classes directly support retail loss prevention: '
            '(1) person = customer counting, occupancy monitoring, unauthorized access detection; '
            '(2) bags (backpack, handbag, suitcase) = concealment risk flagging, entrance bag-check policy enforcement; '
            '(3) high-value items (laptop, cell phone) = checkout verification, electronics display monitoring; '
            '(4) consumables (bottle, cup) = spill hazard detection, inventory gap identification. '
            'ROI gate enables zone-focused monitoring: entrance (count customers + detect bags), '
            'checkout (verify cart items), high-value displays (alert on multiple people), '
            'exits (cross-reference with receipts). Real-time alerts and CSV analytics support loss prevention teams.'
        )
        context_run.font.size = Pt(10)
    
    # Update Git Repository placeholder with red highlighting
    for para in doc.paragraphs:
        if 'Git Repository:' in para.text and 'Your CMP stugit URL here' in para.text:
            para.clear()
            para.add_run('Git Repository: ').bold = True
            placeholder = para.add_run('[REPLACE WITH YOUR STUGIT URL BEFORE SUBMISSION]\n')
            placeholder.font.italic = True
            placeholder.font.color.rgb = RGBColor(255, 0, 0)
            para.add_run('Example: https://stugit.cmp.uea.ac.uk/yourname/webai\n').font.size = Pt(9)
            para.add_run('Clone: ').bold = True
            para.add_run('git clone [YOUR_URL]')
            break
    
    # Save
    doc.save('userguide.docx')
    print('✓ Updated userguide.docx successfully')
    print('  - Added retail security class justification')
    print('  - Highlighted git URL placeholder in RED')
    print('')
    print('ACTION REQUIRED: Open userguide.docx and replace the RED placeholder with your actual stugit URL')

if __name__ == '__main__':
    main()
